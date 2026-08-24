from docx import Document


def extract_docx(
    file_path: str
) -> str:

    document = Document(
        file_path
    )

    content = []

    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            content.append(
                paragraph.text.strip()
            )

    for table in document.tables:

        for row in table.rows:

            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            content.append(
                " | ".join(values)
            )

    return "\n".join(
        content
    )